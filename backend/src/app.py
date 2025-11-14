import io
import json
from typing import List, Optional

import fitz
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse



app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def parse_coords_json(coords_json: str) -> List[List[float]]:
    """
    Парсит координаты из JSON-строки.

    Args:
        coords_json (str): JSON-строка вида [[x1, y1], [x2, y2], ...]

    Returns:
        List[List[float]]: Список координат.

    Raises:
        HTTPException: Если JSON некорректен или пуст.
    """
    try:
        data = json.loads(coords_json)
        if not isinstance(data, list):
            raise ValueError("JSON должен быть списком координат")

        coords = []
        for pair in data:
            if isinstance(pair, list) and len(pair) == 2:
                coords.append([float(pair[0]), float(pair[1])])

        if not coords:
            raise ValueError("Список координат пуст")

        return coords[:10]

    except Exception as exc:
        raise ValueError(f"Некорректные координаты: {exc}")

def parse_coords_excel(excel_file: UploadFile) -> List[List[float]]:
    """
    Парсит координаты из Excel-файла.

    Args:
        excel_file (UploadFile): Загруженный Excel-файл с колонками x и y.

    Returns:
        List[List[float]]: Список координат.

    Raises:
        HTTPException: Если Excel не читается или не содержит координат.
    """
    try:
        wb = load_workbook(filename=io.BytesIO(excel_file.file.read()), data_only=True)
        ws = wb.active

        coords = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            x, y = row[:2]
            if x is not None and y is not None:
                coords.append([float(x), float(y)])
            if len(coords) >= 10:
                break

        if not coords:
            raise ValueError("В Excel-файле нет координат")

        return coords

    except Exception as exc:
        raise ValueError(f"Ошибка Excel: {exc}")


def create_plot(coords: List[List[float]]) -> io.BytesIO:
    """
    Создаёт график по координатам с трендовой линией.

    Args:
        coords (List[List[float]]): Список координат [x, y].

    Returns:
        io.BytesIO: Буфер с PNG-изображением графика.
    """
    xs, ys = zip(*coords)

    fig, ax = plt.subplots()
    ax.scatter(xs, ys)

    sorted_pairs = sorted(zip(xs, ys), key=lambda p: p[0])
    xs_s, ys_s = zip(*sorted_pairs)
    ax.plot(xs_s, ys_s, linestyle="-")

    if len(xs) >= 2:
        p = np.polyfit(xs, ys, 1)
        y_fit = np.polyval(p, xs_s)
        ax.plot(xs_s, y_fit, linestyle="--", linewidth=1)
        ax.text(
            0.02,
            0.98,
            f"Тренд: y={p[0]:.3f}x+{p[1]:.3f}",
            transform=ax.transAxes,
            va="top",
        )

    ax.set_xlabel("X")
    ax.set_ylabel("Y")
    ax.grid(True)

    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


def pdf_to_images(pdf_bytes: bytes) -> List[io.BytesIO]:
    """
    Конвертирует PDF-файл в список изображений страниц.

    Args:
        pdf_bytes (bytes): Содержимое PDF в байтах.

    Returns:
        List[io.BytesIO]: Список буферов с изображениями PNG.
    """
    try:
        if not pdf_bytes or len(pdf_bytes) == 0:
            raise ValueError("PDF-файл пустой")

        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        except RuntimeError as e:
            msg = str(e).lower()

            if "encrypted" in msg or "password" in msg:
                raise ValueError("PDF защищён паролем")
            if "cannot open broken document" in msg:
                raise ValueError("PDF повреждён или имеет неверный формат")
            if "xref" in msg or "no objects" in msg:
                raise ValueError("PDF имеет повреждённую структуру")

            raise ValueError(f"Невозможно открыть PDF: {e}")

        images = []
        for page_number, page in enumerate(doc):
            try:
                pix = page.get_pixmap(dpi=150)
                img = io.BytesIO(pix.tobytes("png"))
                img.seek(0)
                images.append(img)
            except MemoryError:
                raise ValueError(
                    f"Недостаточно памяти для рендеринга страницы {page_number + 1}. "
                    "PDF слишком большой или содержит очень высокое разрешение."
                )
            except Exception as e:
                raise ValueError(f"Ошибка рендеринга страницы {page_number + 1}: {e}")

        return images

    except MemoryError:
        raise ValueError(
            "Недостаточно памяти для обработки PDF. "
            "Попробуйте загрузить файл меньшего размера."
        )
    except Exception as exc:
        raise ValueError(f"Ошибка PDF: {exc}")

@app.post("/generate")
def generate_doc(
        description: Optional[str] = Form(None),
        coords_json: Optional[str] = Form(None),
        excel_file: Optional[UploadFile] = File(None),
        pdf_file: Optional[UploadFile] = File(None),
):
    try:
        if excel_file:
            coords = parse_coords_excel(excel_file)
        elif coords_json:
            coords = parse_coords_json(coords_json)
        else:
            raise ValueError("Не переданы координаты")

        # PDF страницы
        pdf_images = []
        if pdf_file:
            pdf_images = pdf_to_images(pdf_file.file.read())

        # График
        plot_buf = create_plot(coords)

        # DOCX
        doc = Document()
        doc.add_heading("Описание", level=1)
        doc.add_paragraph(description or "Описание работы")

        doc.add_heading("График по координатам", level=1)
        doc.add_paragraph("График.")
        doc.add_picture(plot_buf, width=Inches(6))

        if pdf_images:
            for i, img in enumerate(pdf_images, start=1):
                doc.add_paragraph(f"Страница {i}:")
                doc.add_picture(img, width=Inches(6))

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=generated_document.docx"},
        )

    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail={"error": str(exc)}
        )

    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail={"error": f"Внутренняя ошибка сервера: {exc}"}
        )